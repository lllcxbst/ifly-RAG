import io
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from markdown import markdown
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".html", ".htm"}


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持 {suffix or '无扩展名'} 文件；请上传 TXT、Markdown、PDF、DOCX 或 HTML")

    if suffix in {".txt", ".md", ".markdown"}:
        text = data.decode("utf-8-sig", errors="replace")
        if suffix in {".md", ".markdown"}:
            # Keep markdown headings because the chunker uses them as retrieval context.
            return _clean(text)
        return _clean(text)
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        return _clean("\n\n".join(page.extract_text() or "" for page in reader.pages))
    if suffix == ".docx":
        document = Document(io.BytesIO(data))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if not value:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.removeprefix("Heading ") or "2"
                blocks.append(f"{'#' * int(level)} {value}")
            else:
                blocks.append(value)
        for table in document.tables:
            blocks.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return _clean("\n\n".join(blocks))
    html = data.decode("utf-8-sig", errors="replace")
    if suffix in {".html", ".htm"}:
        soup = BeautifulSoup(html, "html.parser")
    else:
        soup = BeautifulSoup(markdown(html), "html.parser")
    return _clean(soup.get_text("\n"))


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
